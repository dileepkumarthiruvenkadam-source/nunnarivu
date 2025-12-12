import requests
from bs4 import BeautifulSoup
from docx import Document


def scrape_job_details(url: str) -> str:
    """Fetch job description text from a webpage."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, "html.parser")

        # Extract readable text
        paragraphs = soup.find_all(["p", "li"])
        text = "\n".join(p.get_text(strip=True) for p in paragraphs)

        return text if text else "Job details could not be extracted."
    except Exception as e:
        return f"Error fetching job details: {e}"


def create_cover_letter(content: str, output_path: str = "Cover_Letter.docx"):
    """Generate a Word document containing the cover letter."""
    doc = Document()
    doc.add_heading("Cover Letter", level=1)
    doc.add_paragraph(content)
    doc.save(output_path)
    return output_path


def generate_cover_letter(job_url: str, applicant_name: str = "Applicant"):
    """High-level function: scrape → generate text → export to Word."""
    job_text = scrape_job_details(job_url)

    # Very simple template — your Nunnarivu model can improve this later
    cover_letter = f"""
Dear Hiring Team,

I am writing to express my interest in this opportunity. Based on the job posting, here is what I understand about the role:

{job_text}

I believe my skills and background make me a strong fit, and I am excited about the possibility of contributing value.

Best regards,
{applicant_name}
"""

    file_path = create_cover_letter(cover_letter)
    return file_path